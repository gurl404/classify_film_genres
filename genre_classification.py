import pandas as pd
import json
from docx import Document
# Import default_llm information of your choosing

class MovieGenreClassifier:
    @staticmethod
    def classify_genres(project_title: str, synopsis: str) -> dict:
        llm = default_llm
        try:
            llm_response = (
                f"Analyze the following movie:\n"
                f"Title: {project_title}\n"
                f"Synopsis: {synopsis}\n"
                f"Identify all applicable genre tags within this list: Scary, Drama, Fantasy, Comedy." #Add your desired genre classifications here
            )
            response = llm.invoke(llm_response)  
            response_content = response.content

            # Check if the response is empty
            if not response_content.strip():
                print("Received an empty response from the LLM.")
                return {"genres": []}

            # Attempt to parse the response content as JSON
            try:
                recommendations = json.loads(response_content)
            except json.JSONDecodeError:
                # If it's not valid JSON, assume it's a comma-separated string of genres
                genres_list = [genre.strip() for genre in response_content.split(',')]
                recommendations = {"genres": genres_list}

            return recommendations
        except Exception as e:
            print(f"Error calling LLM: {e}. Title: {project_title}, Synopsis: {synopsis}")
            return {"genres": []}

def main():
    # Load the CSV file
    df = pd.read_csv('Film Freeway File')

    # Initialize the classifier
    classifier = MovieGenreClassifier()

    # Create a new Document
    doc = Document()
    doc.add_heading('Movie Genre Classification Results', 0)

    # Iterate through each row in the DataFrame and classify genres
    for index, row in df.iterrows():
        project_title = row['Project Title']
        synopsis = row['Synopsis']
        project_type = row['Project Type']
        existing_genres = row['Genres']

        # Handle cases where project_type is a float (e.g., NaN)
        if isinstance(project_type, float):
            project_type = ""  # Treat NaN as an empty string

        # Skip entries with Project Type "Documentary" or "Music Video"
        if any(ptype.strip() in ["Documentary", "Music Video"] for ptype in project_type.split(',')):
            continue

        # Use existing genres if available
        if pd.notna(existing_genres) and existing_genres.strip():
            genres = {"genres": existing_genres.split(',')}
        else:
            genres = classifier.classify_genres(project_title, synopsis)

        # Add results to the document
        doc.add_heading(f'Project Title: {project_title}', level=1)
        doc.add_paragraph(f'Synopsis: {synopsis}')
        doc.add_paragraph(f'Genres: {", ".join(genres["genres"] if "genres" in genres else [])}')
        doc.add_paragraph('\n')

    doc.save('movie_genre_classification_results.docx')

    print("Results have been saved to 'movie_genre_classification_results.docx'.")

if __name__ == "__main__":
    main()
